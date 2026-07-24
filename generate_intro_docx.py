# -*- coding: utf-8 -*-
"""项目经营范围及优势介绍 - Word 文档"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SIZE_3 = Pt(16)   # 三号
LINE_HEIGHT_28 = Pt(28)


def set_font(run, font_name_cn, size):
    run.font.name = font_name_cn
    run.font.size = size
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), font_name_cn)
    rFonts.set(qn('w:hAnsi'), font_name_cn)
    rFonts.set(qn('w:eastAsia'), font_name_cn)
    rFonts.set(qn('w:cs'), font_name_cn)


def set_line_spacing(paragraph):
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = LINE_HEIGHT_28


def add_h1(doc, text):
    p = doc.add_paragraph()
    p.alignment = 1
    set_line_spacing(p)
    r = p.add_run(text)
    set_font(r, '黑体', SIZE_3)
    r.bold = True
    return p


def add_h2(doc, text):
    p = doc.add_paragraph()
    set_line_spacing(p)
    r = p.add_run(text)
    set_font(r, '楷体', SIZE_3)
    r.bold = True
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(32)
    set_line_spacing(p)
    r = p.add_run(text)
    set_font(r, '仿宋体', SIZE_3)
    return p


doc = Document()

section = doc.sections[0]
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(3.18)
section.right_margin = Cm(3.18)

style = doc.styles['Normal']
style.font.name = '仿宋体'
style.font.size = SIZE_3
rPr = style.element.get_or_add_rPr()
rFonts = rPr.find(qn('w:rFonts'))
if rFonts is None:
    rFonts = OxmlElement('w:rFonts')
    rPr.append(rFonts)
rFonts.set(qn('w:ascii'), '仿宋体')
rFonts.set(qn('w:hAnsi'), '仿宋体')
rFonts.set(qn('w:eastAsia'), '仿宋体')

# 大标题
p = doc.add_paragraph()
p.alignment = 1
set_line_spacing(p)
r = p.add_run('项目经营范围及优势介绍')
set_font(r, '方正小标宋简体', Pt(22))
r.bold = True

doc.add_paragraph()

# 一、经营范围
add_h1(doc, '一、经营范围')
add_body(doc, '本项目以"护花使者"智能植保无人机操控系统为核心，面向油菜花种植产业链提供以下服务：')
add_body(doc, '1. 植保无人机操控软件开发与运营：智能航线规划、AI 参数推荐、实时飞行监控、作业数据管理等移动端应用服务。')
add_body(doc, '2. 油菜病虫害防治技术服务：基于 AI 大模型的病虫害诊断、用药方案推荐、施药时机智能提醒等专业植保咨询。')
add_body(doc, '3. 北斗+GPS 精准定位服务：面向长江流域油菜主产区的亚米级双轨定位及农业气象数据服务。')
add_body(doc, '4. 农业数字化解决方案：为油菜种植合作社、家庭农场、植保服务组织提供定制化的数字植保管理系统。')
add_body(doc, '5. 农业技术培训与推广：飞手操作培训、植保技术普及、田间作业指导等配套服务。')
add_body(doc, '6. 农业物联网设备研发与销售：配套蓝牙通信设备、传感器等智能硬件。')

# 二、核心优势
add_h1(doc, '二、核心优势')

add_h2(doc, '1. 场景深耕优势')
add_body(doc, '专注油菜单一作物，构建从品种、生长阶段到病虫害防治的全链条知识体系，相比通用植保软件更精准、更专业，已形成"油菜植保"垂直领域品牌认知。')

add_h2(doc, '2. AI 技术领先')
add_body(doc, '率先将大语言模型落地到农业植保领域，AI 智能调参准确率高于行业通用方案 20% 以上；新用户 30 分钟即可独立完成专业级作业。')

add_h2(doc, '3. 国产化技术底座')
add_body(doc, '北斗+GPS 双轨定位、自主蓝牙通信协议、国产大模型 API，全栈自主可控，符合国家"信创+乡村振兴"政策导向，享受农机购置补贴支持。')

add_h2(doc, '4. 极致用户体验')
add_body(doc, '极简 3 步操作流程（选区域→AI 调参→起飞），适配大字体、夜间模式等适老化设计，覆盖种植大户、家庭农场、专业服务组织等全类型用户。')

add_h2(doc, '5. 显著经济效益')
add_body(doc, '作业效率提升 30 倍以上，农药利用率从 30% 提升至 60% 以上，亩均增产 8%-15%，农户一季即可收回专业版年费成本。')

add_h2(doc, '6. 产教融合团队')
add_body(doc, '依托湖北职业技术学院，团队兼具软件开发能力与农业技术认知，贴近一线用户需求，迭代速度快，运营成本低。')

add_h2(doc, '7. 完整测试体系')
add_body(doc, '170+ 测试用例全部通过，覆盖单元测试、集成测试、系统测试全流程，产品稳定性和可靠性经过严格验证。')

output_path = r'C:\Users\user\Desktop\护花使者APP\项目经营范围及优势介绍.docx'
doc.save(output_path)
print(f'已生成：{output_path}')
